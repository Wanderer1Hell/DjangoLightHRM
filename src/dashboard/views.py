import os
import calendar as cal
import datetime

from django.core.exceptions import ObjectDoesNotExist
from django.utils import timezone
from django.core.paginator import EmptyPage, PageNotAnInteger, Paginator
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, HttpResponseRedirect
from django.db.models import Q, ExpressionWrapper, When
from django.contrib import messages
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from decimal import Decimal
from dateutil.relativedelta import relativedelta
from django.urls import reverse

from employee.forms import EmployeeCreateForm, EmergencyCreateForm, FamilyCreateForm, BankAccountCreation, \
    DocumentCreateForm, CompanyForm, MilitaryCreateForm, EmploymentHistoryForm
from docx import Document as Docxdoc
from employee.models import *
from employee.forms import LeaveCreationForm
from django.contrib.auth.models import User
import locale

from employee.models import Employee
from hrsuit import settings
from hrsuit.settings import BASE_DIR


# from leave.forms import CommentForm
def dashboard(request):
    '''
	Summary of all apps - display here with charts etc.
	eg.lEAVE - PENDING|APPROVED|RECENT|REJECTED - TOTAL THIS MONTH or NEXT MONTH
	EMPLOYEE - TOTAL | GENDER
	CHART - AVERAGE EMPLOYEE AGES
	'''
    dataset = dict()
    user = request.user

    if not request.user.is_authenticated:
        return redirect('accounts:login')

    employees = Employee.objects.all()
    leaves = Leave.objects.all_pending_leaves()
    employees_birthday = Employee.objects.birthdays_current_month()
    staff_leaves = Leave.objects.filter(employee__user=user)
    dataset['employees'] = employees
    dataset['leaves'] = leaves
    dataset['employees_birthday'] = employees_birthday
    dataset['staff_leaves'] = staff_leaves
    dataset['title'] = 'Панель'

    return render(request, 'dashboard/dashboard_index.html', dataset)


def dashboard_employees(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    dataset = dict()
    departments = Department.objects.all()
    employees = Employee.objects.filter(is_terminated=False)  # Фильтрация по активным сотрудникам
    terminated_employees = Employee.objects.filter(is_terminated=True)  # Фильтрация по не активным сотрудникам

    # pagination
    query = request.GET.get('search')
    if query:
        employees = employees.filter(
            Q(firstname__icontains=query) |
            Q(lastname__icontains=query)
        )

    paginator = Paginator(employees, 10)  # show 10 employee lists per page

    page = request.GET.get('page')
    employees_paginated = paginator.get_page(page)

    dataset['employee_list'] = employees_paginated
    dataset['departments'] = departments
    dataset['all_employees'] = Employee.objects.all_employees()
    dataset['terminated_employees'] = terminated_employees
    dataset['untermintaed_employees'] = employees
    blocked_employees = Employee.objects.all_blocked_employees()

    dataset['blocked_employees'] = blocked_employees
    dataset['title'] = 'Сотрудники'
    return render(request, 'dashboard/employee_app.html', dataset)


def dashboard_employees_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    if request.method == 'POST':
        form = EmployeeCreateForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user

            instance.title = request.POST.get('title')
            instance.image = request.FILES.get('image')
            instance.firstname = request.POST.get('firstname')
            instance.lastname = request.POST.get('lastname')
            instance.othername = request.POST.get('othername')
            instance.sex = request.POST.get('sex')
            instance.bio = request.POST.get('bio')
            instance.birthday = request.POST.get('birthday')

            religion_id = request.POST.get('religion')
            religion = Religion.objects.get(id=religion_id)
            instance.religion = religion

            nationality_id = request.POST.get('nationality')
            nationality = Nationality.objects.get(id=nationality_id)
            instance.nationality = nationality

            department_id = request.POST.get('department')
            department = Department.objects.get(id=department_id)
            instance.department = department

            instance.hometown = request.POST.get('hometown')
            instance.region = request.POST.get('region')
            instance.residence = request.POST.get('residence')
            instance.address = request.POST.get('address')
            instance.education = request.POST.get('education')
            instance.lastwork = request.POST.get('lastwork')
            instance.position = request.POST.get('position')
            instance.ssnitnumber = request.POST.get('ssnitnumber')
            instance.tinnumber = request.POST.get('tinnumber')

            role = request.POST.get('role')
            role_instance = Role.objects.get(id=role)
            instance.role = role_instance

            instance.startdate = request.POST.get('startdate')
            instance.employeetype = request.POST.get('employeetype')
            instance.employeeid = request.POST.get('employeeid')
            instance.dateissued = request.POST.get('dateissued')

            instance.save()

            return redirect('dashboard:employees')
        else:
            messages.error(request, 'Ошибка при создании сотрудника')
            return redirect('dashboard:employeecreate')

    dataset = dict()
    form = EmployeeCreateForm()
    dataset['form'] = form
    dataset['title'] = 'Добавление сотрудника'
    return render(request, 'dashboard/employee_create.html', dataset)


def employee_edit_data(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')
    employee = get_object_or_404(Employee, id=id)
    if request.method == 'POST':
        form = EmployeeCreateForm(request.POST or None, request.FILES or None, instance=employee)
        if form.is_valid():
            instance = form.save(commit=False)

            user = request.POST.get('user')
            assigned_user = User.objects.get(id=user)

            instance.user = assigned_user

            instance.lastname = request.POST.get('lastname')
            instance.firstname = request.POST.get('firstname')
            instance.othername = request.POST.get('othername')
            instance.sex = request.POST.get('sex')
            instance.birthday = request.POST.get('birthday')
            instance.hometown = request.POST.get('hometown')
            instance.title = request.POST.get('title')
            instance.Document_series = request.POST.get('Document_series')
            instance.Document_number = request.POST.get('Document_number')
            instance.Date_issue = request.POST.get('Date_issue')
            instance.Issued = request.POST.get('Issued')
            instance.Unit_code = request.POST.get('Unit_code')
            instance.Date_Registrations = request.POST.get('Date_Registrations')
            instance.Zip_Code = request.POST.get('Zip_Code')

            region_value = request.POST.get('Region')
            if region_value:
                if not region_value.startswith('обл.'):
                    region_value = 'обл.' + region_value
            instance.Region = region_value

            district_value = request.POST.get('District')
            if district_value:
                if not district_value.startswith('р-н.'):
                    district_value = 'р-н.' + district_value
            instance.District = district_value

            settlement_value = request.POST.get('Settlement')
            if settlement_value:
                if not settlement_value.startswith('нп.'):
                    settlement_value = 'нп.' + settlement_value
                instance.Settlement = settlement_value

            street_value = request.POST.get('Street')
            if street_value:
                if not street_value.startswith('ул.'):
                    street_value = 'ул.' + street_value
            instance.Street = street_value

            home_value = request.POST.get('Home')
            if home_value:
                if not home_value.startswith('д.'):
                    home_value = 'д.' + home_value
            instance.Home = home_value

            instance.Corps = request.POST.get('Corps')

            apartment_value = request.POST.get('Apartment')
            if apartment_value:
                if not apartment_value.startswith('кв.'):
                    apartment_value = 'кв.' + apartment_value
            instance.Apartment = apartment_value

            instance.residence = request.POST.get('residence')
            religion_id = request.POST.get('religion')
            religion = Religion.objects.get(id=religion_id)
            instance.religion = religion

            nationality_id = request.POST.get('nationality')
            nationality = Nationality.objects.get(id=nationality_id)
            instance.nationality = nationality

            instance.Resolution = request.POST.get('Resolution')
            instance.ssnitnumber = request.POST.get('ssnitnumber')
            instance.tinnumber = request.POST.get('tinnumber')
            instance.tel = request.POST.get('tel')
            instance.email = request.POST.get('email')

            instance.education = request.POST.get('education')
            instance.lastwork = request.POST.get('lastwork')
            instance.position = request.POST.get('position')

            department_id = request.POST.get('department')
            department = Department.objects.get(id=department_id)
            instance.department = department

            role = request.POST.get('role')
            role_instance = Role.objects.get(id=role)
            instance.role = role_instance

            instance.startdate = request.POST.get('startdate')
            instance.employeetype = request.POST.get('employeetype')
            instance.employeeid = request.POST.get('employeeid')
            instance.dateissued = request.POST.get('dateissued')
            if 'image' in request.FILES:
                instance.image = request.FILES['image']
            instance.bio = request.POST.get('bio')

            # now = datetime.datetime.now()
            # instance.created = now
            # instance.updated = now

            instance.save()
            messages.success(request, 'Аккаунт успешно обновлен',
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:employees')

        else:
            print(form.errors)
            messages.error(request, 'Ошибка обновления учетной записи',
                           extra_tags='alert alert-warning alert-dismissible show')
            return HttpResponse("Данная форма недействительна")

    dataset = dict()
    form = EmployeeCreateForm(request.POST or None, request.FILES or None, instance=employee)
    dataset['form'] = form
    dataset['title'] = 'редактировать - {0}'.format(employee.get_full_name)
    return render(request, 'dashboard/employee_create.html', dataset)


def dashboard_employee_info(request, id):
    if not request.user.is_authenticated:
        return redirect('/')

    employee = get_object_or_404(Employee, id=id)
    employee_emergency_instance = Emergency.objects.filter(employee=employee).first()
    employee_family_instance = Relationship.objects.filter(employee=employee).first()
    employee_military_instance = MilitaryRecord.objects.filter(employee=employee).first()
    bank_instance = Bank.objects.filter(employee=employee).first()
    documents = Document.objects.filter(employee=employee)

    dataset = dict()
    dataset['employee'] = employee
    dataset['emergency'] = employee_emergency_instance
    dataset['family'] = employee_family_instance
    dataset['bank'] = bank_instance
    dataset['documents'] = documents
    dataset['military'] = employee_military_instance
    dataset['title'] = 'Профиль - {0}'.format(employee.get_full_name)
    return render(request, 'dashboard/employee_detail.html', dataset)


def employee_terminate(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    employee = get_object_or_404(Employee, id=id)
    employee.is_terminated = True
    employee.save()

    messages.success(request, 'Сотрудник успешно уволен', extra_tags='alert alert-success alert-dismissible show')

    # Редирект на страницу со всеми сотрудниками
    return redirect('dashboard:employees')


def employee_unterminate(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    employee = get_object_or_404(Employee, id=id)
    employee.is_terminated = False
    employee.save()

    messages.success(request, 'Сотрудник успешно восстановлен', extra_tags='alert alert-success alert-dismissible show')

    # Редирект на страницу со всеми сотрудниками
    return redirect('dashboard:employees')


def dashboard_employee_delete(request, employee_id):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    # Получение объекта сотрудника по его идентификатору
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        # Удаление сотрудника из базы данных
        employee.delete()
        return redirect('dashboard:employees')

    # Отображение страницы подтверждения удаления
    return render(request, 'dashboard/employee_delete.html', {'employee': employee})


# ------------------------- EMERGENCY --------------------------------


def dashboard_emergency_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    emp_name = ''  # Объявление и присвоение значения по умолчанию

    if request.method == 'POST':
        form = EmergencyCreateForm(data=request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            employee_id = request.POST.get('employee')
            employee_object = Employee.objects.get(id=employee_id)
            emp_name = employee_object.get_full_name
            instance.employee = employee_object
            instance.fullname = request.POST.get('fullname')
            instance.tel = request.POST.get('tel')
            instance.location = request.POST.get('location')
            instance.relationship = request.POST.get('relationship')
            instance.save()
            messages.success(request, 'Экстренная информация успешно сохранена для {0}'.format(emp_name),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:emergencycreate')
        else:
            messages.error(request, 'Ошибка при создании экстренной информации {0}'.format(emp_name),
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:emergencycreate')

    dataset = dict()
    form = EmergencyCreateForm()
    dataset['form'] = form
    dataset['title'] = 'Создать экстренный контакт'
    return render(request, 'dashboard/emergency_create.html', dataset)


def dashboard_emergency_update(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser):
        return redirect('/')

    emergency = get_object_or_404(Emergency, id=id)
    employee = emergency.employee
    if request.method == 'POST':
        form = EmergencyCreateForm(data=request.POST, instance=emergency)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.employee = employee
            instance.fullname = request.POST.get('fullname')
            instance.tel = request.POST.get('tel')
            instance.location = request.POST.get('location')
            instance.relationship = request.POST.get('relationship')

            # now = datetime.datetime.now()
            # instance.created = now
            # instance.updated = now

            instance.save()

            messages.success(request, 'Экстренная информация успешно обновлена',
                             extra_tags='alert alert-success alert-dismissible show')
            '''
				NB: redirect() will try to use its given arguments to reverse a URL. 
				This example assumes your URL patterns contain a pattern like this 
				redirect(assumed_url_name,its_assuemed_whatever_instance id)
			'''
            return redirect('dashboard:employeeinfo',
                            id=employee.id)  # worked on redirect to profile and message success and error

    dataset = dict()
    form = EmergencyCreateForm(request.POST or None, instance=emergency)
    dataset['form'] = form
    dataset['title'] = 'Обновление экстренной информации для {0}'.format(employee.get_full_name)

    return render(request, 'dashboard/emergency_create.html', dataset)


# ----------------------------- MILITARY -------------------------------#
def dashboard_military_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    emp_name = ''  # Объявление и присвоение значения по умолчанию

    if request.method == 'POST':
        form = MilitaryCreateForm(data=request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            employee_id = request.POST.get('employee')
            employee_object = Employee.objects.get(id=employee_id)
            emp_name = employee_object.get_full_name
            instance.employee = employee_object
            instance.fullname = request.POST.get('fullname')
            instance.save()
            messages.success(request, 'Информация о ВУ успешно сохранена для {0}'.format(emp_name),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:militarycreate')
        else:
            messages.error(request, 'Ошибка при сохранении информации о ВУ для {0}'.format(emp_name),
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:militaryupdate')

    dataset = dict()
    form = MilitaryCreateForm()
    dataset['form'] = form
    dataset['title'] = 'Добавить документы ВУ'
    return render(request, 'dashboard/military_create.html', dataset)


def dashboard_military_update(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser):
        return redirect('/')

    military = get_object_or_404(MilitaryRecord, id=id)
    employee = military.employee
    if request.method == 'POST':
        form = MilitaryCreateForm(data=request.POST, instance=military)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.employee = employee
            instance.fullname = request.POST.get('fullname')

            # now = datetime.datetime.now()
            # instance.created = now
            # instance.updated = now

            instance.save()

            messages.success(request, 'Информация о ВУ успешно обновлена',
                             extra_tags='alert alert-success alert-dismissible show')
            '''
				NB: redirect() will try to use its given arguments to reverse a URL. 
				This example assumes your URL patterns contain a pattern like this 
				redirect(assumed_url_name,its_assuemed_whatever_instance id)
			'''
            return redirect('dashboard:employeeinfo',
                            id=employee.id)  # worked on redirect to profile and message success and error

    dataset = dict()
    form = MilitaryCreateForm(request.POST or None, instance=military)
    dataset['form'] = form
    dataset['title'] = 'Обновление информации о ВУ для {0}'.format(employee.get_full_name)
    return render(request, 'dashboard/military_create.html', dataset)


# ----------------------------- MILITARY -------------------------------#

# ----------------------------- FAMILY ---------------------------------#


# YOU ARE HERE ---- creation form for Family
def dashboard_family_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')
    if request.method == 'POST':
        form = FamilyCreateForm(data=request.POST or None)
        if form.is_valid():
            instance = form.save(commit=False)
            employee_id = request.POST.get('employee')
            employee_object = get_object_or_404(Employee, id=employee_id)
            instance.employee = employee_object
            instance.status = request.POST.get('status')
            instance.spouse = request.POST.get('spouse')
            instance.occupation = request.POST.get('occupation')
            instance.tel = request.POST.get('tel')
            instance.children = request.POST.get('children')

            # now = datetime.datetime.now()
            # instance.created = now
            # instance.updated = now

            instance.save()

            messages.success(request, 'Успешно созданно для {0}'.format(employee_object),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:familycreate')
        else:
            messages.error(request, 'Не удалось создать для {0}'.format(employee_object),
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:familycreate')

    dataset = dict()

    form = FamilyCreateForm()

    dataset['form'] = form
    dataset['title'] = 'Личная информация'
    return render(request, 'dashboard/family_create_form.html', dataset)


# HERE FAMILY EDIT
def dashboard_family_edit(request, id):
    if not (request.user.is_authenticated and request.user.is_authenticated):
        return redirect('/')
    relation = get_object_or_404(Relationship, id=id)
    employee = relation.employee

    # submitted form - bound form
    if request.method == 'POST':
        form = FamilyCreateForm(data=request.POST, instance=relation)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.employee = employee
            instance.status = request.POST.get('status')
            instance.spouse = request.POST.get('spouse')
            instance.occupation = request.POST.get('occupation')
            instance.tel = request.POST.get('tel')
            instance.children = request.POST.get('children')

            # Recently added 29/03/19
            instance.nextofkin = request.POST.get('nextofkin')
            instance.relationship = request.POST.get('relationship')

            # now = datetime.datetime.now()
            # instance.created = now
            # instance.updated = now

            instance.save()

            messages.success(request, 'Данные, успешно обновленные для {0}'.format(employee.get_full_name),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:familycreate')

        else:
            messages.error(request, 'Не удалось обновить данные для{0}'.format(employee.get_full_name),
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:familycreate')

    dataset = dict()

    form = FamilyCreateForm(request.POST or None, instance=relation)

    dataset['form'] = form
    dataset['title'] = 'Состав семьи'
    return render(request, 'dashboard/family_create_form.html', dataset)


# BANK

def dashboard_bank_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    if request.method == 'POST':
        form = BankAccountCreation(data=request.POST)

        if form.is_valid():
            instance = form.save(commit=False)
            employee_id = request.POST.get('employee')
            employee_object = get_object_or_404(Employee, id=employee_id)
            instance.employee = employee_object
            instance.save()

            messages.success(request, 'Данные успешно созданы для {0}'.format(employee_object.get_full_name),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:bankaccountcreate')
        else:
            messages.error(request, 'Ошибка при создании данных',
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:bankaccountcreate')

    dataset = dict()
    form = BankAccountCreation()
    dataset['form'] = form
    dataset['title'] = 'Банковские данные'
    return render(request, 'dashboard/bank_account_create_form.html', dataset)


def employee_bank_account_update(request, id):
    if not (request.user.is_superuser and request.user.is_authenticated):
        return redirect('/')

    bank_instance_obj = get_object_or_404(Bank, id=id)
    employee = bank_instance_obj.employee
    employee_id = employee.id  # Получаем ID сотрудника
    employment_history_list = EmploymentHistory.objects.filter(employee=employee)
    bank = Bank.objects.get(employee=employee)  # Получаем связанную запись Bank для сотрудника
    bank_id = bank.id

    if request.method == 'POST':
        form = BankAccountCreation(request.POST, instance=bank_instance_obj)

        if form.is_valid():
            instance = form.save(commit=False)
            instance.employee = employee
            instance.save()

            messages.success(request, 'Данные успешно отредактированы для {0}'.format(employee.get_full_name),
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:accountedit', id=bank_id)
        else:
            messages.error(request, 'Ошибка обновления данных',
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:accountedit', id=bank_id)

    dataset = dict()
    form = BankAccountCreation(instance=bank_instance_obj)
    dataset['form'] = form
    dataset['employee_id'] = employee_id  # Передаем employee_id в контексте
    dataset[
        'employment_history_list'] = employment_history_list  # Передаем список записей истории занятости в контексте
    dataset['title'] = 'Данные об условиях труда (заработная плата)'
    return render(request, 'dashboard/bank_account_create_form.html', dataset)


def dashboard_document_create(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    if request.method == 'POST':
        form = DocumentCreateForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            document = request.POST.get('employee')
            document_object = get_object_or_404(Employee, id=document)

            instance.employee = document_object
            instance.save()

            messages.success(request, 'Документ успешно создан',
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:documentcreate')
        else:
            messages.error(request, 'Ошибка при создании документа',
                           extra_tags='alert alert-warning alert-dismissible show')
            return redirect('dashboard:documentcreate')

    form = DocumentCreateForm()
    documents = Document.objects.all()
    dataset = {
        'form': form,
        'title': 'Добавить документ',
        'documents': documents,
    }
    return render(request, 'dashboard/document_create_form.html', dataset)


def dashboard_document_delete(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    document = get_object_or_404(Document, id=id)
    userid = document.employee_id

    if request.method == 'POST':
        # Удаление файла
        document_path = document.document_file.path
        if os.path.exists(document_path):
            os.remove(document_path)

        # Удаление записи из БД
        document.delete()

        messages.success(request, 'Документ успешно удален',
                         extra_tags='alert alert-success alert-dismissible show')
        return redirect('dashboard:employeeinfo', id=userid)
    else:
        dataset = {
            'title': 'Удаление документа',
            'document': document,
        }
        return render(request, 'dashboard/document_delete_form.html', dataset)


def dashboard_document_edit(request, id):
    pass


# ---------------------LEAVE-------------------------------------------


def leave_creation(request):
    if not request.user.is_authenticated:
        return redirect('accounts:login')
    if request.method == 'POST':
        form = LeaveCreationForm(data=request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            user = request.user
            instance.user = user
            instance.save()

            # print(instance.defaultdays)
            messages.success(request, 'Отпуск сохранён',
                             extra_tags='alert alert-success alert-dismissible show')
            return redirect('dashboard:createleave')

        messages.error(request, 'Не получилось сохранить отпуск, пожалуйста, проверьте даты',
                       extra_tags='alert alert-warning alert-dismissible show')
        return redirect('dashboard:createleave')

    dataset = dict()
    form = LeaveCreationForm()
    dataset['form'] = form
    dataset['title'] = 'Сохранить отпуск сотрудника'
    return render(request, 'dashboard/create_leave.html', dataset)


# return HttpResponse('leave creation form')


# def leave_creation(request):


# 	if request.method == 'POST':
# 		form = LeaveCreationForm(data = request.POST)
# 		cform = CommentForm(data = request.POST)
# 		if form.is_valid() and cform.is_valid():
# 			instance = form.save(commit = False)
# 			user = request.user
# 			instance.user = user
# 			instance.save()
# 			print(instance)

# 			# Commment form save  logic
# 			comment_inst = cform.save(commit = False)
# 			# comment_inst.leave = instance
# 			# comment_inst.comment = request.POST['comment']
# 			cinstance.save()

# 			return HttpResponse('success')

# 		else:
# 			return HttpResponse('error')


# 	dataset = dict()

# 	form = LeaveCreationForm()
# 	cform = CommentForm()
# 	dataset['form'] = form
# 	dataset['cform'] = cform
# 	return render(request,'dashboard/create_leave.html',dataset)


def leaves_list(request):
    if not (request.user.is_staff and request.user.is_superuser):
        return redirect('/')
    leaves = Leave.objects.all_pending_leaves()
    return render(request, 'dashboard/leaves_recent.html', {'leave_list': leaves, 'title': 'Список отпусков'})


def leaves_approved_list(request):
    if not (request.user.is_superuser and request.user.is_staff):
        return redirect('/')
    leaves = Leave.objects.all_approved_leaves()  # approved leaves -> calling model manager method
    return render(request, 'dashboard/leaves_approved.html', {'leave_list': leaves, 'title': 'Утверждённые отпуска'})


def leaves_view(request, id):
    if not request.user.is_authenticated:
        return redirect('/')

    leave = get_object_or_404(Leave, id=id)
    employee = leave.employee
    return render(request, 'dashboard/leave_detail_view.html', {'leave': leave, 'employee': employee,
                                                                'title': f'{employee.user.username}-{leave.status} leave'})


def approve_leave(request, id):
    if not (request.user.is_superuser and request.user.is_authenticated):
        return redirect('/')
    leave = get_object_or_404(Leave, id=id)
    user = leave.employee.user
    employee = Employee.objects.filter(user=user).first()  # используйте first() вместо [0]
    leave.approve_leave()

    messages.error(request, 'Отпуск успешно утвержден для {0}'.format(employee.get_full_name()),
                   extra_tags='alert alert-success alert-dismissible show')
    return redirect('dashboard:userleaveview', id=id)


def cancel_leaves_list(request):
    if not (request.user.is_superuser and request.user.is_authenticated):
        return redirect('/')
    leaves = Leave.objects.all_cancel_leaves()
    return render(request, 'dashboard/leaves_cancel.html', {'leave_list_cancel': leaves, 'title': 'Отменённые отпуска'})


def unapprove_leave(request, id):
    if not (request.user.is_authenticated and request.user.is_superuser):
        return redirect('/')
    leave = get_object_or_404(Leave, id=id)
    leave.unapprove_leave
    return redirect('dashboard:leaveslist')  # redirect to unapproved list


def cancel_leave(request, id):
    if not (request.user.is_superuser and request.user.is_authenticated):
        return redirect('/')
    leave = get_object_or_404(Leave, id=id)
    leave.leaves_cancel

    messages.success(request, 'Заявление на отпуск отменено', extra_tags='alert alert-success alert-dismissible show')
    return redirect('dashboard:canceleaveslist')  # work on redirecting to instance leave - detail view


# Current section -> here
def uncancel_leave(request, id):
    if not (request.user.is_superuser and request.user.is_authenticated):
        return redirect('/')
    leave = get_object_or_404(Leave, id=id)
    leave.status = 'pending'
    leave.is_approved = False
    leave.save()
    messages.success(request, 'Заявление на отпуск, сейчас находится в списке ожидающих рассмотрения.',
                     extra_tags='alert alert-success alert-dismissible show')
    return redirect('dashboard:canceleaveslist')  # work on redirecting to instance leave - detail view


def leave_rejected_list(request):
    dataset = dict()
    leave = Leave.objects.all_rejected_leaves()

    dataset['leave_list_rejected'] = leave
    return render(request, 'dashboard/rejected_leaves_list.html', dataset)


def reject_leave(request, id):
    dataset = dict()
    leave = get_object_or_404(Leave, id=id)
    leave.reject_leave
    messages.success(request, 'Заявление на отпуск отклонено', extra_tags='alert alert-success alert-dismissible show')
    return redirect('dashboard:leavesrejected')


# return HttpResponse(id)


def unreject_leave(request, id):
    leave = get_object_or_404(Leave, id=id)
    leave.status = 'pending'
    leave.is_approved = False
    leave.save()
    messages.success(request, 'Заявление на отпуск теперь находится в списке ожидающих рассмотрения ',
                     extra_tags='alert alert-success alert-dismissible show')

    return redirect('dashboard:leavesrejected')


def view_my_leave_table(request):
    # work on the logics
    if request.user.is_authenticated:
        user = request.user
        leaves = Leave.objects.filter(user=user)
        employee = Employee.objects.filter(user=user).first()
        dataset = dict()
        dataset['leave_list'] = leaves
        dataset['employee'] = employee
        dataset['title'] = 'Отпуска'
    else:
        return redirect('accounts:login')
    return render(request, 'dashboard/staff_leaves_table.html', dataset)


# Birthday
def birthday_this_month(request):
    if not request.user.is_authenticated:
        return redirect('/')

    employees = Employee.objects.birthdays_current_month()
    locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
    month = datetime.date.today().strftime(
        '%B').capitalize()  # am using this to get the month for template rendering- making it dynamic
    context = {
        'birthdays': employees,
        'month': month,
        'count_birthdays': employees.count(),
        'title': 'Дни рождения'
    }
    return render(request, 'dashboard/birthdays_this_month.html', context)


# Organizational data
def company_list(request):
    companies = Company.objects.all()
    return render(request, 'dashboard/company_list.html', {'companies': companies})


def company_create_or_update(request, id=None):
    company = None
    if id:
        company = get_object_or_404(Company, id=id)

    if request.method == 'POST':
        form = CompanyForm(request.POST, instance=company)
        if form.is_valid():
            form.save()
            return redirect('dashboard:companylist')
    else:
        form = CompanyForm(instance=company)

    return render(request, 'dashboard/company_create_or_update.html', {'form': form})


def company_delete(request, id):
    company = get_object_or_404(Company, id=id)
    if request.method == 'POST':
        company.delete()
        return redirect('dashboard:companylist')
    return render(request, 'dashboard/company_delete.html', {'company': company})


def fill_template_t8(request):
    if request.method == 'POST' and 'generate-document-t8' in request.POST:

        employee_id = request.POST.get('employee-t8')
        employee = Employee.objects.get(id=employee_id)

        role_id = employee.role_id
        role = Role.objects.get(id=role_id)
        role_title = role.name

        departament_id = employee.department_id
        departament = Department.objects.get(id=departament_id)
        departament_title = departament.name

        company = Company.objects.first()

        condition = request.POST.get('condition-t8')

        template_path = os.path.join(os.path.dirname(BASE_DIR), 'src', 'static_cdn', 'doctemplates', 't8.docx')

        document_number = request.POST.get('document-number-t8')
        date_compiled = request.POST.get('date-compiled-t8')
        data = {
            '{{name}}': employee.firstname,
            '{{lastname}}': employee.lastname,
            '{{othername}}': employee.othername,
            '{{role}}': role_title,
            '{{departament}}': departament_title,
            '{{company}}': company.name,
            '{{position_master}}': company.position,
            '{{employeeid}}': employee.employeeid,
            '{{document_number}}': document_number,
            '{{date_compiled}}': date_compiled,
            '{{condition}}': condition,
        }

        filled_doc = fill_template_document(template_path, data)

        # Генерация пути к файлу документа
        folder_path = os.path.join('documents', timezone.now().strftime('%Y/%m/%d'))
        os.makedirs(os.path.join(settings.MEDIA_ROOT, folder_path), exist_ok=True)
        file_name = 'prikaz_o_uvolnenii.docx'

        # Проверка наличия файла с таким же именем
        i = 1
        while os.path.exists(os.path.join(settings.MEDIA_ROOT, folder_path, file_name)):
            # Добавление уникального номера к имени файла
            base_name, extension = os.path.splitext(file_name)
            file_name = f'{base_name}_{i}{extension}'
            i += 1

        # Сохранение заполненного документа
        file_path = os.path.join(folder_path, file_name)
        filled_doc.save(os.path.join(settings.MEDIA_ROOT, file_path))

        # Создание и сохранение объекта Document
        document = Document.objects.create(
            employee=employee,
            document_file=file_path,
            filename='Приказ о прекращении трудового договора с работником'
        )

        # Отправка заполненного документа в качестве ответа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        filled_doc.save(response)

        return response
    # Если метод запроса не POST или кнопка "Сгенерировать документ" не была нажата,
    # просто рендерим шаблон без выполнения генерации документа
    employees = Employee.objects.all()
    return render(request, 'fill_template.html', {'employees': employees})


def fill_template_t6(request):
    if request.method == 'POST' and 'generate-document-t6' in request.POST:
        employee_id = request.POST.get('employee-t6')
        employee = Employee.objects.get(id=employee_id)

        role_id = employee.role_id
        role = Role.objects.get(id=role_id)
        role_title = role.name

        departament_id = employee.department_id
        departament = Department.objects.get(id=departament_id)
        departament_title = departament.name

        company = Company.objects.first()

        leavetype = employee.leave_set.first().leavetype if employee.leave_set.exists() else ' '
        leave_days = employee.leave_set.first().leave_days if employee.leave_set.exists() else 10

        template_path = os.path.join(os.path.dirname(BASE_DIR), 'src', 'static_cdn', 'doctemplates', 't6.docx')

        document_number = request.POST.get('document-number-t6')
        date_compiled = request.POST.get('date-compiled-t6')
        data = {
            '{{name}}': employee.firstname,
            '{{lastname}}': employee.lastname,
            '{{othername}}': employee.othername,
            '{{role}}': role_title,
            '{{departament}}': departament_title,
            '{{company}}': company.name,
            '{{position_master}}': company.position,
            '{{employeeid}}': employee.employeeid,
            '{{document_number}}': document_number,
            '{{date_compiled}}': date_compiled,
            '{{leavetype}}': leavetype,
            '{{leave_days}}': leave_days,
        }

        filled_doc = fill_template_document(template_path, data)

        # Генерация пути к файлу документа
        folder_path = os.path.join('documents', timezone.now().strftime('%Y/%m/%d'))
        os.makedirs(os.path.join(settings.MEDIA_ROOT, folder_path), exist_ok=True)
        file_name = 'prikaz_na_otpusk.docx'

        # Проверка наличия файла с таким же именем
        i = 1
        while os.path.exists(os.path.join(settings.MEDIA_ROOT, folder_path, file_name)):
            # Добавление уникального номера к имени файла
            base_name, extension = os.path.splitext(file_name)
            file_name = f'{base_name}_{i}{extension}'
            i += 1

        # Сохранение заполненного документа
        file_path = os.path.join(folder_path, file_name)
        filled_doc.save(os.path.join(settings.MEDIA_ROOT, file_path))

        # Создание и сохранение объекта Document
        document = Document.objects.create(
            employee=employee,
            document_file=file_path,
            filename='Приказ на отпуск'
        )

        # Отправка заполненного документа в качестве ответа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        filled_doc.save(response)

        return response

    # Если метод запроса не POST или кнопка "Сгенерировать документ" не была нажата,
    # просто рендерим шаблон без выполнения генерации документа
    employees = Employee.objects.all()
    return render(request, 'fill_template.html', {'employees': employees})


def fill_template_t11(request):
    if request.method == 'POST' and 'generate-document-t11' in request.POST:
        employee_id = request.POST.get('employee-t11')
        employee = Employee.objects.get(id=employee_id)

        role_id = employee.role_id
        role = Role.objects.get(id=role_id)
        role_title = role.name

        departament_id = employee.department_id
        departament = Department.objects.get(id=departament_id)
        departament_title = departament.name

        company = Company.objects.first()
        condition = request.POST.get('condition-t11')
        motive = request.POST.get('motive-t11')
        amount = (request.POST.get('amount-t11'))
        amountwr = request.POST.get('amountwr-t11')

        template_path = os.path.join(os.path.dirname(BASE_DIR), 'src', 'static_cdn', 'doctemplates', 't11.docx')

        document_number = request.POST.get('document-number-t11')
        date_compiled = request.POST.get('date-compiled-t11')
        data = {
            '{{name}}': employee.firstname,
            '{{lastname}}': employee.lastname,
            '{{othername}}': employee.othername,
            '{{role}}': role_title,
            '{{departament}}': departament_title,
            '{{company}}': company.name,
            '{{position_master}}': company.position,
            '{{employeeid}}': employee.employeeid,
            '{{document_number}}': document_number,
            '{{date_compiled}}': date_compiled,
            '{{condition}}': condition,
            '{{motive}}': motive,
            '{{amount}}': amount,
            '{{amountwr}}': amountwr,

        }

        filled_doc = fill_template_document(template_path, data)

        # Генерация пути к файлу документа
        folder_path = os.path.join('documents', timezone.now().strftime('%Y/%m/%d'))
        os.makedirs(os.path.join(settings.MEDIA_ROOT, folder_path), exist_ok=True)
        file_name = 'prikaz-o-pooschrenii.docx'

        # Проверка наличия файла с таким же именем
        i = 1
        while os.path.exists(os.path.join(settings.MEDIA_ROOT, folder_path, file_name)):
            # Добавление уникального номера к имени файла
            base_name, extension = os.path.splitext(file_name)
            file_name = f'{base_name}_{i}{extension}'
            i += 1

        # Сохранение заполненного документа
        file_path = os.path.join(folder_path, file_name)
        filled_doc.save(os.path.join(settings.MEDIA_ROOT, file_path))

        # Создание и сохранение объекта Document
        document = Document.objects.create(
            employee=employee,
            document_file=file_path,
            filename='Приказ о поощрении работника'
        )

        # Отправка заполненного документа в качестве ответа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        filled_doc.save(response)

        return response

    # Если метод запроса не POST или кнопка "Сгенерировать документ" не была нажата,
    # просто рендерим шаблон без выполнения генерации документа
    employees = Employee.objects.all()
    return render(request, 'fill_template.html', {'employees': employees})


def fill_template_t2(request):
    if request.method == 'POST' and 'generate-document-t2' in request.POST:
        employee_id = request.POST.get('employee-t2')
        employee = Employee.objects.get(id=employee_id)

        company = Company.objects.first()
        condition = request.POST.get('condition-t2')
        date_incident = request.POST.get('date_incident-t2')
        reason = request.POST.get('reason-t2')

        template_path = os.path.join(os.path.dirname(BASE_DIR), 'src', 'static_cdn', 'doctemplates', 't2.docx')

        document_number = request.POST.get('document-number-t2')
        date_compiled = request.POST.get('date-compiled-t2')
        data = {
            '{{name}}': employee.firstname,
            '{{lastname}}': employee.lastname,
            '{{othername}}': employee.othername,
            '{{company}}': company.name,
            '{{condition}}': condition,
            '{{document_number}}': document_number,
            '{{date_compiled}}': date_compiled,
            '{{date_incident}}': date_incident,
            '{{reason}}': reason,
        }

        filled_doc = fill_template_document(template_path, data)

        # Генерация пути к файлу документа
        folder_path = os.path.join('documents', timezone.now().strftime('%Y/%m/%d'))
        os.makedirs(os.path.join(settings.MEDIA_ROOT, folder_path), exist_ok=True)
        file_name = 'prikaza-o-vzyskanii.docx'

        # Проверка наличия файла с таким же именем
        i = 1
        while os.path.exists(os.path.join(settings.MEDIA_ROOT, folder_path, file_name)):
            # Добавление уникального номера к имени файла
            base_name, extension = os.path.splitext(file_name)
            file_name = f'{base_name}_{i}{extension}'
            i += 1

        # Сохранение заполненного документа
        file_path = os.path.join(folder_path, file_name)
        filled_doc.save(os.path.join(settings.MEDIA_ROOT, file_path))

        # Создание и сохранение объекта Document
        document = Document.objects.create(
            employee=employee,
            document_file=file_path,
            filename='Приказ о взыскании'
        )

        # Отправка заполненного документа в качестве ответа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        filled_doc.save(response)

        return response

    # Если метод запроса не POST или кнопка "Сгенерировать документ" не была нажата,
    # просто рендерим шаблон без выполнения генерации документа
    employees = Employee.objects.all()
    return render(request, 'fill_template.html', {'employees': employees})


def fill_template(request):
    if request.method == 'POST' and 'generate-document' in request.POST:
        employee_id = request.POST.get('employee')
        employee = Employee.objects.get(id=employee_id)

        role_id = employee.role_id
        role = Role.objects.get(id=role_id)
        role_title = role.name

        departament_id = employee.department_id
        departament = Department.objects.get(id=departament_id)
        departament_title = departament.name

        bank = Bank.objects.get(employee_id=employee.id)  # Получение связанной банковской записи
        salary = bank.salary
        rubles = int(salary)
        kopeks = int((salary - rubles) * 100)

        company = Company.objects.first()

        template_path = os.path.join(os.path.dirname(BASE_DIR), 'src', 'static_cdn', 'doctemplates', 't1.docx')

        document_number = request.POST.get('document-number')
        date_compiled = request.POST.get('date-compiled')
        condition = request.POST.get('condition')
        data = {
            '{{name}}': employee.firstname,
            '{{lastname}}': employee.lastname,
            '{{othername}}': employee.othername,
            '{{role}}': role_title,
            '{{departament}}': departament_title,
            '{{company}}': company.name,
            '{{position_master}}': company.position,
            '{{salary_rubles}}': rubles,
            '{{salary_kopeks}}': kopeks,
            '{{employeeid}}': employee.employeeid,
            '{{document_number}}': document_number,
            '{{date_compiled}}': date_compiled,
            '{{condition}}': condition,
        }

        filled_doc = fill_template_document(template_path, data)

        # Генерация пути к файлу документа
        folder_path = os.path.join('documents', timezone.now().strftime('%Y/%m/%d'))
        os.makedirs(os.path.join(settings.MEDIA_ROOT, folder_path), exist_ok=True)
        file_name = 'prikaz_o_prieme_na_rabotu.docx'

        # Проверка наличия файла с таким же именем
        i = 1
        while os.path.exists(os.path.join(settings.MEDIA_ROOT, folder_path, file_name)):
            # Добавление уникального номера к имени файла
            base_name, extension = os.path.splitext(file_name)
            file_name = f'{base_name}_{i}{extension}'
            i += 1

        # Сохранение заполненного документа
        file_path = os.path.join(folder_path, file_name)
        filled_doc.save(os.path.join(settings.MEDIA_ROOT, file_path))

        # Создание и сохранение объекта Document
        document = Document.objects.create(
            employee=employee,
            document_file=file_path,
            filename='Приказ о приёме на работу'
        )

        # Отправка заполненного документа в качестве ответа
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        filled_doc.save(response)

        return response
    # Если метод запроса не POST или кнопка "Сгенерировать документ" не была нажата,
    # просто рендерим шаблон без выполнения генерации документа
    employees = Employee.objects.all()
    return render(request, 'fill_template.html', {'employees': employees})


def fill_template_document(template_path, data):
    doc = Docxdoc(template_path)

    filled_style = doc.styles.add_style('FilledData', WD_STYLE_TYPE.PARAGRAPH)
    filled_style.font.name = 'Times New Roman'
    filled_style.font.size = Pt(10)
    filled_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                paragraph.style = filled_style

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                        cell.paragraphs[0].style = filled_style

    return doc


def employee_info_schedule(request):
    if not (request.user.is_authenticated and request.user.is_superuser and request.user.is_staff):
        return redirect('/')

    dataset = dict()
    departments = Department.objects.all()
    employees = Employee.objects.filter(is_terminated=False)  # Фильтрация по активным сотрудникам
    terminated_employees = Employee.objects.filter(is_terminated=True)  # Фильтрация по не активным сотрудникам

    # pagination
    query = request.GET.get('search')
    if query:
        employees = employees.filter(
            Q(firstname__icontains=query) |
            Q(lastname__icontains=query)
        )

    paginator = Paginator(employees, 10)  # show 10 employee lists per page

    page = request.GET.get('page')
    employees_paginated = paginator.get_page(page)

    dataset['employee_list'] = employees_paginated
    dataset['departments'] = departments
    dataset['all_employees'] = Employee.objects.all_employees()
    dataset['terminated_employees'] = terminated_employees
    dataset['untermintaed_employees'] = employees
    blocked_employees = Employee.objects.all_blocked_employees()

    dataset['blocked_employees'] = blocked_employees
    dataset['title'] = 'Расписание'
    return render(request, 'dashboard/employee_info_schedule.html', dataset)


def holiday_request(request):
    if request.method == 'POST':
        date = request.POST.get('date')
        name = request.POST.get('name')

        Holiday.objects.create(date=date, name=name)

        return redirect('dashboard:holiday_request')

    holidays = Holiday.objects.all()

    context = {
        'holidays': holidays,
    }
    return render(request, 'dashboard/holiday_request.html', context)


def holiday_delete(request, holiday_id):
    holiday = get_object_or_404(Holiday, id=holiday_id)
    holiday.delete()
    return redirect('dashboard:holiday_request')


def work_schedule(request, employee_id):
    employee = Employee.objects.get(id=employee_id)
    bank_records = Bank.objects.filter(employee=employee)

    work_schedule = []

    status_colors = {
        'П': '#FFD700',  # Золотистый цвет для статуса 'П'
        '0': '#FFB6C1',  # Розовый цвет для статуса '0'
        '8': '#90EE90',  # Светло-зеленый цвет для статуса '8'
        '2.5': '#90EE90',  # Светло-зеленый цвет для статуса '2.5'
    }

    for record in bank_records:
        if record.work_schedule:
            work_schedule.append(record.work_schedule)

    today = datetime.date.today()
    num_days = cal.monthrange(today.year, today.month)[1]  # Количество дней в текущем месяце

    schedule = []

    if bank_records and bank_records[0].work_schedule == 1:
        for i in range(num_days):
            status = '8'
            if i % 7 in (5, 6):
                status = '0'
            elif Holiday.objects.filter(date=datetime.date(today.year, today.month, i + 1)).exists():
                status = 'П'
            if status in status_colors:
                status_color = status_colors[status]
            else:
                status_color = '#FFFFFF'  # Белый цвет для других статусов
            schedule.append({'status': status, 'color': status_color})
    elif bank_records and bank_records[0].work_schedule == 2:
        weekdays = [0, 1, 2, 3]
        for i in range(num_days):
            if i % 7 in (5, 6):
                status = '0'
            elif i % 7 in weekdays:
                if Holiday.objects.filter(date=datetime.date(today.year, today.month, i + 1)).exists():
                    status = 'П'
                else:
                    status = '2.5'
            else:
                status = '0'
            if status in status_colors:
                status_color = status_colors[status]
            else:
                status_color = '#FFFFFF'  # Белый цвет для других статусов
            schedule.append({'status': status, 'color': status_color})
    elif bank_records and bank_records[0].work_schedule == 3:
        for i in range(num_days):
            if i % 4 < 2:
                if Holiday.objects.filter(date=datetime.date(today.year, today.month, i + 1)).exists():
                    status = 'П'
                else:
                    status = '8'
            else:
                status = '0'
            if status in status_colors:
                status_color = status_colors[status]
            else:
                status_color = '#FFFFFF'  # Белый цвет для других статусов
            schedule.append({'status': status, 'color': status_color})
    else:
        schedule = [{'status': 'none', 'color': '#FFFFFF'}] * num_days

    context = {
        'employee': employee,
        'today': today,
        'num_days': num_days,
        'schedule': schedule,
    }

    return render(request, 'dashboard/work_schedule.html', context)


# ---------------------experience -------------------------------------------

def employment_history_add(request, employee_id):
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        form = EmploymentHistoryForm(request.POST)
        if form.is_valid():
            employment_history = form.save(commit=False)
            employment_history.employee = employee
            employment_history.save()

            # Получение объекта Bank по employee_id
            bank = Bank.objects.get(employee_id=employee_id)
            bank_id = bank.id

            messages.success(request, 'Запись успешно добавлена',
                             extra_tags='alert alert-success alert-dismissible show')

            return redirect('dashboard:accountedit', id=bank_id)
    else:
        form = EmploymentHistoryForm()

    context = {
        'form': form,
        'title': 'Добавить данные',
        'employee_id': employee_id,
        'employment_history_list': EmploymentHistory.objects.filter(employee=employee)
    }

    return render(request, 'dashboard/employment_history_add.html', context)


def employment_history_edit(request, employment_history_id):
    employment_history = get_object_or_404(EmploymentHistory, id=employment_history_id)

    if request.method == 'POST':
        form = EmploymentHistoryForm(request.POST, instance=employment_history)
        if form.is_valid():
            form.save()
            messages.success(request, 'Запись успешно обновлена',
                             extra_tags='alert alert-success alert-dismissible show')

            return redirect('dashboard:employment_history_edit', employment_history_id=employment_history.id)

    else:
        form = EmploymentHistoryForm(instance=employment_history)

    context = {
        'form': form,
        'title': 'Редактировать запись',
        'employment_history': employment_history,
        'employment_history_list': EmploymentHistory.objects.all()
    }
    return render(request, 'dashboard/employment_history_edit.html', context)


def calculate_total_experience(employee):
    employment_history = EmploymentHistory.objects.filter(employee=employee)
    total_experience = relativedelta()

    for history in employment_history:
        if history.experience_start_date and history.experience_end_date:
            experience_start_date = history.experience_start_date
            experience_end_date = history.experience_end_date
            total_experience += relativedelta(experience_end_date, experience_start_date)

    return total_experience


def calculate_payment(total_experience, mrot, salary, start_date, end_date):
    total_payment = Decimal(0)

    if total_experience.years < 1 and total_experience.months < 6:  # Менее 6 месяцев стажа
        if mrot is None:
            raise ValueError("Введите МРОТ")

        daily_payment = mrot * Decimal('0.6') / Decimal('365')  # Один день больничного стоит 60% от МРОТ
    else:
        employment_years = total_experience.years  # Количество полных лет работы
        employment_months = total_experience.months  # Количество полных месяцев работы
        total_experience_months = employment_years * 12 + employment_months  # Общее количество полных месяцев работы

        if total_experience_months < 60:  # От 6 месяцев до 5 лет стажа
            daily_payment = salary * Decimal('0.6') / Decimal('365')  # Один день больничного стоит 60% от оклада
        elif total_experience_months < 96:  # От 5 до 8 лет стажа
            daily_payment = salary * Decimal('0.8') / Decimal('365')  # Один день больничного стоит 80% от оклада
        else:  # Свыше 8 лет стажа
            daily_payment = salary / Decimal('365')  # Один день больничного стоит 100% от оклада

    total_days = (end_date - start_date).days + 1  # Количество дней в больничном периоде
    total_payment = daily_payment * Decimal(total_days)

    return total_payment


def sick_leave(request):
    if request.method == 'POST':
        employee_id = request.POST.get('employee')
        start_date_str = request.POST.get('start_date')
        end_date_str = request.POST.get('end_date')
        diagnosis = request.POST.get('diagnosis')
        conclusion = request.POST.get('conclusion')
        issuing_institution = request.POST.get('issuing_institution')
        doctor_name = request.POST.get('doctor_name')

        if not start_date_str or not end_date_str:
            error_message = "Введите дату"
            employees = Employee.objects.all()
            sick_leaves = SickLeave.objects.all()
            return render(request, 'dashboard/sick_leave.html',
                          {'employees': employees, 'sick_leaves': sick_leaves, 'error_message': error_message})

        employee = Employee.objects.get(id=employee_id)

        try:
            start_date = datetime.datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except ValueError:
            error_message = "Неправильный формат даты"
            employees = Employee.objects.all()
            sick_leaves = SickLeave.objects.all()
            return render(request, 'dashboard/sick_leave.html',
                          {'employees': employees, 'sick_leaves': sick_leaves, 'error_message': error_message})

        bank = Bank.objects.get(employee=employee)
        salary = bank.salary

        total_experience = calculate_total_experience(employee)

        if total_experience.years < 1 and total_experience.months < 6:
            mrot_str = request.POST.get('mrot')
            if not mrot_str:
                error_message = "Введите МРОТ"
                employees = Employee.objects.all()
                sick_leaves = SickLeave.objects.all()
                return render(request, 'dashboard/sick_leave.html',
                              {'employees': employees, 'sick_leaves': sick_leaves, 'error_message': error_message})

            try:
                mrot = Decimal(mrot_str)
            except ValueError:
                error_message = "Неправильный формат значения МРОТ"
                employees = Employee.objects.all()
                sick_leaves = SickLeave.objects.all()
                return render(request, 'dashboard/sick_leave.html',
                              {'employees': employees, 'sick_leaves': sick_leaves, 'error_message': error_message})
        else:
            mrot = 0

        total_payment = calculate_payment(total_experience, mrot, salary, start_date, end_date)
        sick_leave = SickLeave.objects.create(
            employee=employee,
            start_date=start_date,
            end_date=end_date,
            diagnosis=diagnosis,
            conclusion=conclusion,
            issuing_institution=issuing_institution,
            doctor_name=doctor_name,
            payment=total_payment
        )

    elif request.method == 'POST' and 'delete' in request.POST:
        sick_leave_id = request.POST.get('delete')
        sick_leave = SickLeave.objects.get(id=sick_leave_id)
        sick_leave.delete()
        return redirect('sick_leave')

    employees = Employee.objects.all()
    sick_leaves = SickLeave.objects.all()
    error_message = request.POST.get('error_message')

    return render(request, 'dashboard/sick_leave.html',
                  {'employees': employees, 'sick_leaves': sick_leaves, 'error_message': error_message})


def delete_sick_leave(request, sick_leave_id):
    sick_leave = get_object_or_404(SickLeave, pk=sick_leave_id)
    sick_leave.delete()
    return redirect('dashboard:sick_leave')
