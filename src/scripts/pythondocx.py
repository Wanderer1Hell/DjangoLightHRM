from docx import Document

def fill_template(template_path, data):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    return doc

# Пример использования
template_path = 't1.docx'  # Путь к шаблонному документу
data = {
    '{{name}}': 'Anna',
    '{{email}}': 'anna@duck.com'
}

filled_doc = fill_template(template_path, data)
filled_doc.save('output.docx')  # Сохранение заполненного документа
